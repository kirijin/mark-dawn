#!/usr/bin/env python3
"""Build a high-quality .docx from markdown text.

Uses python-docx for programmatic control — fonts, spacing, styles,
tables, lists, page numbers. Falls back to pandoc --reference-doc
if python-docx is unavailable.
"""
import re, sys, shutil, subprocess, tempfile
from pathlib import Path

# ── python-docx import (graceful fallback) ──
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYDOCX = True
except ImportError:
    HAS_PYDOCX = False

# ── Defaults (user-overridable via env) ──
FONT     = "Calibri"
FONT_SIZE = 11       # pt
MARGIN   = 2.54      # cm (1 inch)
LINE_SP  = 1.15
HEADING_FONTS = {"Heading 1": 18, "Heading 2": 15, "Heading 3": 13}

def markdown_to_docx(md_text: str, output_path: str | Path,
                      title: str = "Converted Document") -> Path:
    """Main entry: parse markdown, write high-quality docx."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # Fallback: pandoc
        return _pandoc_fallback(md_text, output_path, title)

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin    = Cm(MARGIN)
    section.bottom_margin = Cm(MARGIN)
    section.left_margin   = Cm(MARGIN)
    section.right_margin  = Cm(MARGIN)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT
    font.size = Pt(FONT_SIZE)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)  # near-black
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SP

    # ── Heading styles ──
    for level, size in HEADING_FONTS.items():
        s = doc.styles[level]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)  # dark blue
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(8)
        s.paragraph_format.line_spacing = 1.15

    # ── List Bullet style ──
    lb = doc.styles['List Bullet']
    lb.font.name = FONT
    lb.font.size = Pt(FONT_SIZE)
    lb.paragraph_format.space_after = Pt(3)
    lb.paragraph_format.left_indent = Cm(1.27)

    # ── Code style ──
    cs = doc.styles.add_style('CodeBlock', 1)  # paragraph style
    cs.font.name = 'Consolas'
    cs.font.size = Pt(9)
    cs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    cs.paragraph_format.space_before = Pt(4)
    cs.paragraph_format.space_after = Pt(4)
    cs.paragraph_format.left_indent = Cm(0.5)
    # Gray shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    shd.set(qn('w:val'), 'clear')
    cs.element.get_or_add_pPr().append(shd)

    # ── Page numbers ──
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

    # ── Add title as document heading ──
    if title:
        doc.add_heading(title, level=0)

    # ── Parse markdown ──
    _build_from_md(doc, md_text)

    doc.save(str(output_path))
    return output_path

def _build_from_md(doc, md_text):
    """Parse markdown lines and populate the Document."""
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code block ──
        if stripped.startswith('```'):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
                i += 1
                continue
            else:
                in_code = True
                i += 1
                continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Skip empty lines (just paragraph separators) ──
        if not stripped:
            i += 1
            continue

        # ── Headings ──
        hm = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2)
            _add_heading(doc, text, level)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^[-*_]{3,}$', stripped):
            doc.add_paragraph('—' * 20)
            i += 1
            continue

        # ── Bullet list ──
        if re.match(r'^[\s]*[-*+]\s+', stripped):
            text = re.sub(r'^[\s]*[-*+]\s+', '', stripped)
            _add_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        # ── Numbered list ──
        nm = re.match(r'^[\s]*\d+[.)]\s+(.+)$', stripped)
        if nm:
            text = nm.group(1)
            _add_paragraph(doc, text, style='List Number')
            i += 1
            continue

        # ── Table (simple pipe table) ──
        if '|' in stripped and stripped.count('|') >= 2:
            table_lines = _gather_table_lines(lines, i)
            if len(table_lines) >= 2:  # header + at least one row
                _add_table(doc, table_lines)
                i += len(table_lines)
                continue

        # ── Regular paragraph ──
        _add_paragraph(doc, stripped)
        i += 1

def _add_heading(doc, text, level):
    """Add a heading with inline formatting."""
    h = doc.add_heading('', level=min(level, 3))
    _add_inline_run(h, text)
    return h

def _add_paragraph(doc, text, style=None):
    """Add paragraph with bold/italic/inline-code support."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _add_inline_run(p, text)
    return p

def _add_inline_run(para, text):
    """Parse inline markdown (**bold**, *italic*, `code`) into runs."""
    # Tokenise: split on **, *, ` markers
    pattern = r'(\*\*.*?\*\*|[*].*?[*]|`.*?`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            para.add_run(part)

def _add_code_block(doc, lines):
    """Dedented code block lines as styled paragraphs."""
    if not lines:
        return
    min_indent = _min_indent(lines)
    for line in lines:
        text = line[min_indent:] if len(line) > min_indent else ''
        doc.add_paragraph(text, style='CodeBlock')

def _min_indent(lines):
    non_empty = [l for l in lines if l.strip()]
    if not non_empty:
        return 0
    spaces = [len(l) - len(l.lstrip()) for l in non_empty]
    return min(spaces)

def _gather_table_lines(lines, start):
    """Collect consecutive pipe-table lines from start."""
    result = []
    for i in range(start, min(start + 200, len(lines))):
        l = lines[i].strip()
        if not l or l.startswith('```') or re.match(r'^#{1,6}\s', l):
            break
        if '|' in l and l.count('|') >= 2:
            result.append(l)
        elif result:
            break  # non-table line after table started
        else:
            break
    return result

def _add_table(doc, table_lines):
    """Build a Word table from pipe-table markdown."""
    # Skip separator rows (| --- | --- |)
    data_rows = [l for l in table_lines if not re.match(r'^[\s|:-]+$', l)]
    if len(data_rows) < 2:
        return
    # Parse cells
    rows = []
    for line in data_rows:
        cells = [c.strip() for c in line.split('|')]
        # Remove leading/trailing empty (from leading/trailing |)
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    if ncols < 2:
        return
    # Normalise width
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = 1  # CENTER

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(10)
            if ri == 0:
                run.bold = True  # header row

def _pandoc_fallback(md_text, output_path, title):
    """Convert via pandoc with a reference template if available."""
    with tempfile.TemporaryDirectory() as td:
        md_path = Path(td) / "input.md"
        md_path.write_text(md_text, encoding="utf-8")

        # Try with a built-in reference
        cmd = ["pandoc", str(md_path), "-o", str(output_path),
               "--metadata", f"title={title}",
               "-V", f"mainfont={FONT}"]
        subprocess.run(cmd, capture_output=True, text=True, timeout=120, check=True)

    return output_path

if __name__ == "__main__":
    # CLI: docx_styler.py input.md output.docx
    if len(sys.argv) < 3:
        print("Usage: docx_styler.py INPUT.md OUTPUT.docx", file=sys.stderr)
        sys.exit(1)
    md = Path(sys.argv[1]).read_text(encoding="utf-8")
    out = markdown_to_docx(md, sys.argv[2])
    print(f"OK: {out}")
